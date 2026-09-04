# src/m09_docs.py
"""
Modul für Dokumenten-Management und RAG-Ingestion.
Behandelt Upload, Text-Extraktion, Chunking und Embedding von Dokumenten.
"""
import os
import hashlib
import json
import math
import threading
from pathlib import Path
from datetime import datetime, timezone
from typing import List, Optional, Tuple, Any, NamedTuple
from sqlmodel import select, Session

# PDF-Handling: Versuche verschiedene Bibliotheken
try:
    import pypdf as PyPDF2  # pypdf ist der Nachfolger von PyPDF2, API-kompatibel
except ImportError:
    PyPDF2 = None

try:
    import pdfplumber
except ImportError:
    pdfplumber = None

try:
    import pandas as pd
except ImportError:
    pd = None

try:
    from docx import Document as DocxDocument
except ImportError:
    DocxDocument = None

from .m01_config import get_settings
from .m03_db import (
    get_session, Document, DocumentChunk, ProjectDocumentLink, 
    DOC_LIMITS, DOCUMENT_CLASSIFICATIONS
)
from .m09_rag import embed_text, embed_texts_batch, EMBEDDING_MODEL, clear_rag_cache

S = get_settings()

# Verzeichnis für hochgeladene Dokumente
DOCS_DIR = Path(S.data_dir) / "rag" / "docs"
DOCS_DIR.mkdir(parents=True, exist_ok=True)

MIN_PDF_TEXT_CHARS = 80
VISION_OCR_MAX_PAGES = 12
VISION_OCR_PAGE_PROMPT = (
    "Extrahiere den vollständigen lesbaren Text dieser Dokumentenseite (Deutsch/Englisch). "
    "Tabellen zeilenweise, Listen als Aufzählung. Nur Dokumenttext — keine Beschreibung des Bildes."
)


class IngestResult(NamedTuple):
    ok: bool
    message: str
    status: str = "ok"

    def __iter__(self):
        yield self.ok
        yield self.message


def _seed_keywords_for_document(doc_id: int, provider: str = "openai", model: str = "gpt-4o-mini") -> None:
    """
    Generiert retrieval_keywords für alle Chunks eines Dokuments (Background-Thread).
    Überspringt Chunks die bereits Keywords haben (idempotent).
    """
    import logging
    log = logging.getLogger(__name__)
    try:
        from .m08_llm import generate_chunk_keywords
        with get_session() as ses:
            chunks = ses.exec(
                select(DocumentChunk)
                .where(DocumentChunk.document_id == doc_id)
                .where(DocumentChunk.retrieval_keywords == None)
            ).all()

        for chunk in chunks:
            keywords = generate_chunk_keywords(chunk.chunk_text or "", provider=provider, model=model)
            if keywords:
                with get_session() as ses:
                    db_chunk = ses.get(DocumentChunk, chunk.id)
                    if db_chunk:
                        db_chunk.retrieval_keywords = json.dumps(keywords, ensure_ascii=False)
                        ses.add(db_chunk)
                        ses.commit()
        log.info("retrieval_keywords geseedet für doc_id=%s (%d Chunks)", doc_id, len(chunks))
    except Exception as exc:
        logging.getLogger(__name__).warning("Keyword-Seeding fehlgeschlagen für doc_id=%s: %s", doc_id, exc)


def calculate_sha256(file_bytes: bytes) -> str:
    """Berechnet den SHA256-Hash von Bytes."""
    sha256_hash = hashlib.sha256()
    sha256_hash.update(file_bytes)
    return sha256_hash.hexdigest()


def _slice_pdf_pages(
    pages,
    *,
    max_pages: int | None = None,
    page_start: int = 1,
    page_end: int | None = None,
):
    """Seitenauswahl: 1-basiert, page_end inklusive."""
    total = len(pages)
    if page_start > 1 or page_end is not None:
        start = max(0, min(page_start - 1, total))
        end = min(page_end or total, total)
        if end < start:
            end = start
        return pages[start:end]
    if max_pages is not None:
        return pages[:max_pages]
    return pages


def extract_text_from_pdf(
    file_path: Path,
    max_pages: int | None = None,
    *,
    page_start: int = 1,
    page_end: int | None = None,
) -> str:
    """Extrahiert Text aus einer PDF-Datei mit Fallback-Strategien."""
    text_content = []
    
    # Strategie 1: PyPDF2
    if PyPDF2:
        try:
            with open(file_path, "rb") as f:
                reader = PyPDF2.PdfReader(f)
                pages = _slice_pdf_pages(
                    reader.pages,
                    max_pages=max_pages,
                    page_start=page_start,
                    page_end=page_end,
                )
                for page in pages:
                    text = page.extract_text()
                    if text:
                        text_content.append(text)
            
            full_text = "\n".join(text_content)
            if full_text.strip():
                return full_text
        except Exception as e:
            print(f"PyPDF2 extraction failed for {file_path}: {e}")
    
    # Strategie 2: pdfplumber (besser für Layouts)
    if pdfplumber:
        try:
            text_content = []
            with pdfplumber.open(file_path) as pdf:
                pages = _slice_pdf_pages(
                    pdf.pages,
                    max_pages=max_pages,
                    page_start=page_start,
                    page_end=page_end,
                )
                for page in pages:
                    text = page.extract_text()
                    if text:
                        text_content.append(text)
            
            full_text = "\n".join(text_content)
            if full_text.strip():
                return full_text
        except Exception as e:
            print(f"pdfplumber extraction failed for {file_path}: {e}")

    return ""


def _pdf_pages_as_png_bytes(file_path: Path, max_pages: int) -> list[tuple[bytes, str]]:
    from pdf2image import convert_from_path
    from io import BytesIO

    images = convert_from_path(
        str(file_path),
        dpi=150,
        first_page=1,
        last_page=max(1, max_pages),
    )
    out: list[tuple[bytes, str]] = []
    for im in images[:max_pages]:
        buf = BytesIO()
        im.save(buf, format="PNG")
        out.append((buf.getvalue(), "image/png"))
    return out


def _extract_pdf_text_via_vision(file_path: Path, *, max_pages: int = VISION_OCR_MAX_PAGES) -> str:
    """Optionaler Fallback: Vision-LLM liest gescannte PDF-Seiten (wenn ocrmypdf nicht reicht)."""
    from .m08_llm import have_key, try_models_with_messages
    from .m17_visual_lab_refs import resolve_vision_provider_model

    provider, model = resolve_vision_provider_model("", "")
    if not have_key(provider):
        return ""
    try:
        pages = _pdf_pages_as_png_bytes(file_path, max_pages)
    except Exception:
        return ""
    if not pages:
        return ""

    parts: list[str] = []
    for i, (img_bytes, mime) in enumerate(pages, 1):
        raw = try_models_with_messages(
            provider,
            "Du extrahierst Text aus gescannten Dokumentseiten für eine Volltext-Suche.",
            [{"role": "user", "content": f"Seite {i}:\n{VISION_OCR_PAGE_PROMPT}"}],
            max_tokens=2500,
            temperature=0.0,
            model=model,
            images=[(img_bytes, mime)],
        )
        if raw and raw.strip():
            parts.append(raw.strip())
    return "\n\n".join(parts)


def extract_pdf_text_with_fallback(
    file_path: Path,
    *,
    max_pages: int | None = None,
    use_vision_fallback: bool = True,
) -> tuple[str, dict[str, Any]]:
    """
    PDF-Text: native → ocrmypdf → optional Vision-LLM.
    Gibt (text, meta) zurück; meta für UI/Warnungen.
    """
    from .m11_vision import apply_ocr_to_pdf, is_pdf_scanned, pdf_extracted_text_is_sparse

    meta: dict[str, Any] = {
        "scanned_detected": False,
        "ocr_attempted": False,
        "ocr_success": False,
        "vision_attempted": False,
        "vision_success": False,
        "extraction_method": "native",
    }

    text = extract_text_from_pdf(file_path, max_pages=max_pages)
    sparse = pdf_extracted_text_is_sparse(text, min_chars=MIN_PDF_TEXT_CHARS)
    if sparse:
        meta["scanned_detected"] = is_pdf_scanned(str(file_path)) or not (text or "").strip()

    if sparse:
        meta["ocr_attempted"] = True
        if apply_ocr_to_pdf(str(file_path)):
            ocr_text = extract_text_from_pdf(file_path, max_pages=max_pages)
            if not pdf_extracted_text_is_sparse(ocr_text, min_chars=MIN_PDF_TEXT_CHARS):
                text = ocr_text
                meta["ocr_success"] = True
                meta["extraction_method"] = "ocrmypdf"

    if use_vision_fallback and pdf_extracted_text_is_sparse(text, min_chars=MIN_PDF_TEXT_CHARS):
        meta["vision_attempted"] = True
        vision_text = _extract_pdf_text_via_vision(file_path)
        if not pdf_extracted_text_is_sparse(vision_text, min_chars=MIN_PDF_TEXT_CHARS):
            text = vision_text
            meta["vision_success"] = True
            meta["extraction_method"] = "vision_llm"

    return text or "", meta


def extract_text_from_docx(file_path: Path) -> str:
    """Extrahiert Text aus einer Word (.docx) Datei."""
    if not file_path.exists():
        return ""
    
    if not DocxDocument:
        return ""  # python-docx nicht installiert
    
    try:
        doc = DocxDocument(file_path)
        text = "\n".join([paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip()])
        return text
    except Exception as e:
        print(f"docx extraction failed for {file_path}: {e}")
        return ""


def chunk_text(text: str, chunk_size: int = 1000, overlap: int = 200) -> List[str]:
    """
    Teilt Text in überlappende Chunks auf.
    Einfache Implementierung basierend auf Zeichenlänge.
    """
    if not text:
        return []
    
    if chunk_size is None:
        chunk_size = 1000
    
    # Smart overlap calculation if not provided or default
    if overlap == 200:
        overlap = max(chunk_size // 5, 50)
    
    chunks = []
    start = 0
    text_len = len(text)
    
    while start < text_len:
        end = start + chunk_size
        chunk = text[start:end]
        
        # Versuche, an einem Satzende oder Zeilenumbruch zu schneiden, wenn möglich
        if end < text_len:
            last_period = chunk.rfind('.')
            last_newline = chunk.rfind('\n')
            cut_point = max(last_period, last_newline)
            
            if cut_point > chunk_size * 0.5:  # Nur schneiden, wenn wir nicht zu viel verlieren
                end = start + cut_point + 1
                chunk = text[start:end]
        
        chunks.append(chunk)
        start = end - overlap
        
    return chunks


def process_csv_to_chunks(file_path: Path, delimiter: str = ";") -> Tuple[bool, str, List[dict]]:
    """
    Verarbeitet CSV-Datei für Batch-QA:
    - Jede Zeile wird ein Chunk (als JSON gespeichert)
    - Validiert Pflicht-Spalten: Nr, Lieferant, Frage
    - Returns: (success, message, chunks_as_dicts)
    """
    if pd is None:
        return False, "pandas ist nicht installiert - CSV-Import nicht möglich", []
    
    try:
        df = pd.read_csv(file_path, sep=delimiter, encoding="utf-8")
    except Exception as e:
        try:
            df = pd.read_csv(file_path, sep=delimiter, encoding="latin-1")
        except Exception as e2:
            return False, f"CSV konnte nicht gelesen werden: {e2}", []
    
    # Validierung: Pflicht-Spalten (flexibel mit Punkt-Varianten)
    # Normalisiere Spaltennamen: entferne Punkte und Leerzeichen, lowercase
    df_normalized_cols = {col.strip().rstrip('.').lower(): col for col in df.columns}
    
    required_cols = ["nr", "lieferant", "frage"]
    missing = []
    col_mapping = {}  # Maps: required_name -> actual_column_name
    
    for req_col in required_cols:
        if req_col in df_normalized_cols:
            col_mapping[req_col] = df_normalized_cols[req_col]
        else:
            missing.append(req_col.capitalize())
    
    if missing:
        return False, f"CSV fehlen Pflicht-Spalten: {', '.join(missing)}. Gefunden: {list(df.columns)}", []
    
    # Stelle sicher dass "Antwort"-Spalte existiert (optional im Input, aber wir brauchen sie für Output)
    antwort_col = df_normalized_cols.get("antwort")
    if not antwort_col:
        df["Antwort"] = ""
        antwort_col = "Antwort"
    
    chunks = []
    for idx, row in df.iterrows():
        chunk_dict = {
            "Nr": str(row.get(col_mapping["nr"], idx)),
            "Lieferant": str(row.get(col_mapping["lieferant"], "")),
            "Frage": str(row.get(col_mapping["frage"], "")),
            "Antwort": str(row.get(antwort_col, ""))
        }
        chunks.append(chunk_dict)
    
    return True, f"CSV erfolgreich verarbeitet: {len(chunks)} Zeilen/Fragen", chunks


def process_generic_csv_rows(file_path: Path, delimiter: str = ";") -> Tuple[bool, str, List[dict]]:
    """FAQ/Listen-CSV: jede Zeile ein Chunk (ohne Batch-QA-Pflichtspalten)."""
    if pd is None:
        return False, "pandas ist nicht installiert", []
    try:
        df = pd.read_csv(file_path, sep=delimiter, encoding="utf-8")
    except Exception:
        try:
            df = pd.read_csv(file_path, sep=delimiter, encoding="latin-1")
        except Exception as exc:
            return False, f"CSV konnte nicht gelesen werden: {exc}", []
    rows: list[dict] = []
    for idx, row in df.iterrows():
        parts = [f"{col}: {row[col]}" for col in df.columns if str(row[col]).strip()]
        text = " | ".join(parts)
        if text.strip():
            rows.append({"Nr": str(idx), "Lieferant": "", "Frage": text, "Antwort": ""})
    if not rows:
        return False, "CSV enthält keine Datenzeilen", []
    return True, f"CSV zeilenweise: {len(rows)} Zeilen", rows


def process_xlsx_to_rows(file_path: Path) -> Tuple[bool, str, List[dict]]:
    """Excel-Tabellen: jede Zeile als strukturierter Chunk (Preisblatt etc.)."""
    if pd is None:
        return False, "pandas ist nicht installiert — XLSX nicht unterstützt", []
    try:
        df = pd.read_excel(file_path)
    except Exception as exc:
        return False, f"XLSX konnte nicht gelesen werden: {exc}", []
    df = df.fillna("")
    rows: list[dict] = []
    for idx, row in df.iterrows():
        row_dict = {str(col).strip(): str(row[col]).strip() for col in df.columns}
        if any(v for v in row_dict.values()):
            rows.append(row_dict)
    if not rows:
        return False, "XLSX enthält keine Datenzeilen", []
    return True, f"XLSX: {len(rows)} Zeilen", rows


def _embed_and_store_chunks(
    session,
    doc: Document,
    file_name: str,
    classification: str,
    doc_subtype: str | None,
    chunk_size: int,
    *,
    csv_chunks: list[dict] | None = None,
    text_content: str = "",
) -> Tuple[bool, str]:
    """Gemeinsame Chunk/Embed-Logik für ingest und force_rechunk."""
    old_chunks = session.exec(
        select(DocumentChunk).where(DocumentChunk.document_id == doc.id)
    ).all()
    for ch in old_chunks:
        session.delete(ch)
    session.flush()

    if csv_chunks:
        doc.chunk_count = len(csv_chunks)
        doc.chunk_size_used = 0
        session.add(doc)
        frage_texts = []
        for row in csv_chunks:
            if "Frage" in row:
                nr = row.get("Nr", "?")
                frage = row.get("Frage", "")
                prefix = f"[CSV | {file_name} | Frage {nr}]\n"
                frage_texts.append(prefix + str(frage))
            else:
                prefix = chunk_meta_prefix(classification, file_name, doc_subtype)
                frage_texts.append(prefix + json.dumps(row, ensure_ascii=False))
        embeddings = embed_texts_batch(frage_texts)
        for i, row_dict in enumerate(csv_chunks):
            if "Frage" in row_dict:
                nr = row_dict.get("Nr", "?")
                contextual_chunk = f"[CSV | {file_name} | Frage {nr}]\n{json.dumps(row_dict, ensure_ascii=False)}"
            else:
                prefix = chunk_meta_prefix(classification, file_name, doc_subtype)
                contextual_chunk = prefix + json.dumps(row_dict, ensure_ascii=False)
            emb = embeddings[i] if i < len(embeddings) else None
            session.add(DocumentChunk(
                document_id=doc.id, chunk_index=i, chunk_text=contextual_chunk,
                embedding=json.dumps(emb) if emb else None,
                embedding_model=EMBEDDING_MODEL,
                tokens_count=len(contextual_chunk) // 4,
            ))
        session.commit()
        clear_rag_cache()
        threading.Thread(target=_seed_keywords_for_document, args=(doc.id,), daemon=True).start()
        return True, f"{doc.chunk_count} Tabellenzeilen indexiert"
    if text_content:
        chunks = chunk_text(text_content, chunk_size=chunk_size)
        doc.chunk_count = len(chunks)
        doc.chunk_size_used = chunk_size
        session.add(doc)
        prefix = chunk_meta_prefix(classification, file_name, doc_subtype)
        contextual_chunks = [prefix + c for c in chunks]
        embeddings = embed_texts_batch(contextual_chunks)
        for i, contextual_chunk in enumerate(contextual_chunks):
            emb = embeddings[i] if i < len(embeddings) else None
            session.add(DocumentChunk(
                document_id=doc.id, chunk_index=i, chunk_text=contextual_chunk,
                embedding=json.dumps(emb) if emb else None,
                embedding_model=EMBEDDING_MODEL,
                tokens_count=len(contextual_chunk) // 4,
            ))
        session.commit()
        clear_rag_cache()
        threading.Thread(target=_seed_keywords_for_document, args=(doc.id,), daemon=True).start()
        return True, f"{doc.chunk_count} Chunks (Grösse {chunk_size})"
    return False, "Kein Text extrahiert"


def force_rechunk_document(doc_id: int, chunk_size: int | None = None) -> Tuple[bool, str]:
    """Bestehendes Dokument neu chunken/embedden (gleiche Datei, neue Chunk-Grösse)."""
    with get_session() as session:
        doc = session.get(Document, doc_id)
        if not doc or doc.is_deleted:
            return False, "Dokument nicht gefunden"
        path = Path(doc.file_path)
        if not path.exists():
            return False, "Dateidatei fehlt auf dem Server"
        size = chunk_size or doc.chunk_size_used or 1000
        ext = path.suffix.lower()
        csv_chunks = None
        text_content = ""
        if ext == ".csv":
            ok, msg, data = process_csv_to_chunks(path)
            if not ok:
                ok, msg, data = process_generic_csv_rows(path)
            if not ok:
                return False, msg
            csv_chunks = data
        elif ext in (".xlsx", ".xls"):
            ok, msg, data = process_xlsx_to_rows(path)
            if not ok:
                return False, msg
            csv_chunks = data
        elif ext == ".pdf":
            text_content, _pdf_meta = extract_pdf_text_with_fallback(path)
        elif ext == ".docx":
            text_content = extract_text_from_docx(path)
        elif ext in (".md", ".txt", ".json", ".yaml", ".yml"):
            try:
                text_content = path.read_text(encoding="utf-8")
            except Exception:
                text_content = path.read_text(encoding="latin-1", errors="replace")
        else:
            try:
                text_content = path.read_text(encoding="utf-8", errors="replace")
            except Exception:
                pass
        ok, msg = _embed_and_store_chunks(
            session, doc, doc.filename, doc.classification, doc.doc_subtype, size,
            csv_chunks=csv_chunks, text_content=text_content,
        )
        if ok:
            doc.chunk_size_used = size
            session.add(doc)
            session.commit()
        return ok, msg


def chunk_meta_prefix(
    classification: str,
    file_name: str,
    doc_subtype: str | None = None,
) -> str:
    """Chunk-Prefix für RAG: optionaler Subtyp übersteuert nur die Anzeige, nicht den Dateiname-Boost."""
    if doc_subtype:
        return f"[{classification} · {doc_subtype} | {file_name}]\n"
    return f"[{classification} | {file_name}]\n"


def _ingest_status_from_meta(meta: dict[str, Any], *, has_chunks: bool) -> str:
    if has_chunks:
        if meta.get("vision_success"):
            return "vision_ok"
        if meta.get("ocr_success"):
            return "ocr_ok"
        return "ok"
    if meta.get("ocr_attempted") or meta.get("vision_attempted") or meta.get("scanned_detected"):
        return "ocr_failed"
    return "no_text"


def _extraction_method_label(meta: dict[str, Any]) -> str:
    method = meta.get("extraction_method") or "native"
    if method == "ocrmypdf":
        return " (OCR)"
    if method == "vision_llm":
        return " (Vision-KI)"
    return ""


def ingest_document(
    file_name: str,
    file_bytes: bytes,
    classification: str,
    chunk_size: int = 1000,
    csv_delimiter: str = ";",
    linked_role_keys: list[str] | None = None,
    doc_subtype: str | None = None,
    *,
    use_vision_fallback: bool = True,
) -> IngestResult:
    """
    Verarbeitet ein hochgeladenes Dokument:
    1. Hash prüfen (Duplikate)
    2. Speichern
    3. Text extrahieren
    4. Chunken & Embedden
    5. DB-Einträge erstellen
    
    Für CSV: csv_delimiter bestimmt das Trennzeichen (default: ";")
    linked_role_keys: Optional, für "Pflichtenheft (Rolle)" - Liste von role.key Werten
    doc_subtype: Optional, feinere Art innerhalb der Klassifikation (z.B. Preisblatt)
    use_vision_fallback: Bei gescannten PDFs optional Vision-LLM nach ocrmypdf
    """
    file_hash = calculate_sha256(file_bytes)
    pdf_meta: dict[str, Any] = {}
    
    with get_session() as session:
        # Check Duplikate
        existing = session.exec(select(Document).where(Document.sha256_hash == file_hash)).first()
        if existing:
            if existing.is_deleted:
                # Hard-Reset: alte Chunks löschen, neu chunken & embedden mit neuer chunk_size
                old_chunks = session.exec(
                    select(DocumentChunk).where(DocumentChunk.document_id == existing.id)
                ).all()
                for ch in old_chunks:
                    session.delete(ch)
                session.flush()
                existing.is_deleted = False
                existing.classification = classification
                existing.chunk_size_used = chunk_size
                existing.chunk_count = 0
                existing.doc_subtype = (doc_subtype or "").strip() or None
                if linked_role_keys is not None:
                    existing.linked_role_keys = json.dumps(linked_role_keys)
                session.add(existing)
                session.commit()
                session.refresh(existing)
                # Jetzt neu chunken & embedden (gleicher Code wie Neu-Upload)
                _reingest_doc = existing
            else:
                return IngestResult(
                    False,
                    f"Dokument '{file_name}' existiert bereits (ID: {existing.id}).",
                    "linked",
                )
        else:
            _reingest_doc = None

        # Datei speichern (auch bei Re-Ingest: neue Datei mit neuem Timestamp)
        safe_name = Path(file_name).name
        ts_prefix = datetime.now().strftime("%Y%m%d_%H%M%S")
        saved_filename = f"{ts_prefix}_{safe_name}"
        file_path = DOCS_DIR / saved_filename

        try:
            file_path.write_bytes(file_bytes)
        except Exception as e:
            return IngestResult(False, f"Fehler beim Speichern der Datei: {e}", "error")

        # Text extrahieren
        text_content = ""
        csv_chunks = None  # Für strukturierte CSV-Verarbeitung
        ext = file_path.suffix.lower()

        if ext == ".csv":
            success, message, csv_data = process_csv_to_chunks(file_path, delimiter=csv_delimiter)
            if not success and classification == "FAQ/Fragen-Katalog":
                success, message, csv_data = process_generic_csv_rows(file_path, delimiter=csv_delimiter)
            if not success:
                return IngestResult(False, f"CSV-Verarbeitung fehlgeschlagen: {message}", "error")
            csv_chunks = csv_data
        elif ext in (".xlsx", ".xls"):
            success, message, csv_data = process_xlsx_to_rows(file_path)
            if not success:
                return IngestResult(False, f"XLSX-Verarbeitung fehlgeschlagen: {message}", "error")
            csv_chunks = csv_data
        elif ext == ".pdf":
            text_content, pdf_meta = extract_pdf_text_with_fallback(
                file_path, use_vision_fallback=use_vision_fallback,
            )
        elif ext == ".docx":
            text_content = extract_text_from_docx(file_path)
        elif ext in [".md", ".txt", ".json", ".yaml", ".yml"]:
            try:
                text_content = file_path.read_text(encoding="utf-8")
            except:
                try:
                    text_content = file_path.read_text(encoding="latin-1")
                except:
                    text_content = ""
        else:
            try:
                text_content = file_path.read_text(encoding="utf-8")
            except:
                pass

        # DB Eintrag: bei Re-Ingest bestehenden Eintrag aktualisieren, sonst neu anlegen
        linked_keys_json = json.dumps(linked_role_keys) if linked_role_keys else None

        with get_session() as session2:
            if _reingest_doc is not None:
                doc = session2.get(Document, _reingest_doc.id)
                doc.file_path = str(file_path)
                doc.doc_subtype = (doc_subtype or "").strip() or None
            else:
                doc = Document(
                    filename=file_name,
                    sha256_hash=file_hash,
                    classification=classification,
                    file_path=str(file_path),
                    file_size=len(file_bytes),
                    embedding_model=EMBEDDING_MODEL,
                    chunk_count=0,
                    chunk_size_used=chunk_size,
                    linked_role_keys=linked_keys_json,
                    doc_subtype=(doc_subtype or "").strip() or None,
                )
                session2.add(doc)
            session2.commit()
            session2.refresh(doc)

            ok, msg = _embed_and_store_chunks(
                session2, doc, file_name, classification, doc_subtype, chunk_size,
                csv_chunks=csv_chunks, text_content=text_content,
            )
            if csv_chunks or text_content:
                if ok:
                    label = "neu indexiert" if _reingest_doc is not None else "erfolgreich importiert"
                    suffix = _extraction_method_label(pdf_meta) if ext == ".pdf" else ""
                    status = _ingest_status_from_meta(pdf_meta, has_chunks=True)
                    return IngestResult(
                        True,
                        f"Dokument '{file_name}' {label}{suffix}: {msg}",
                        status,
                    )
                return IngestResult(False, msg, "error")
            status = _ingest_status_from_meta(pdf_meta, has_chunks=False)
            if status == "ocr_failed":
                msg = (
                    f"Dokument '{file_name}' gespeichert — Gescannt: OCR fehlgeschlagen, "
                    "0 Chunks indexiert. Datei manuell als durchsuchbares PDF vorbereiten oder erneut hochladen."
                )
            else:
                msg = f"Dokument '{file_name}' gespeichert, aber kein Text extrahiert."
            return IngestResult(True, msg, status)


def delete_document(doc_id: int) -> bool:
    """Löscht ein Dokument (Soft-Delete)."""
    with get_session() as session:
        doc = session.get(Document, doc_id)
        if not doc:
            return False
        
        doc.is_deleted = True
        session.add(doc)
        session.commit()
        return True

def list_documents(include_deleted: bool = False) -> List[Document]:
    with get_session() as session:
        query = select(Document)
        if not include_deleted:
            query = query.where(Document.is_deleted == False)
        query = query.order_by(Document.uploaded_at.desc())
        return session.exec(query).all()


def get_document_by_id(doc_id: int) -> Optional[Document]:
    with get_session() as session:
        doc = session.get(Document, doc_id)
        if not doc or doc.is_deleted:
            return None
        return doc


def get_document_by_sha256(file_hash: str, include_deleted: bool = False) -> Optional[Document]:
    with get_session() as session:
        doc = session.exec(select(Document).where(Document.sha256_hash == file_hash)).first()
        if not doc:
            return None
        if doc.is_deleted and not include_deleted:
            return None
        return doc


def resolve_document_path(doc: Document) -> Optional[Path]:
    """Absoluter Pfad zur Datei — DB-Pfad oder Fallback data/rag/docs/."""
    candidates: list[Path] = []
    if doc.file_path:
        candidates.append(Path(doc.file_path))
        candidates.append(DOCS_DIR / Path(doc.file_path).name)
    candidates.append(DOCS_DIR / doc.filename)
    seen: set[Path] = set()
    for raw in candidates:
        try:
            p = raw.resolve()
        except OSError:
            continue
        if p in seen:
            continue
        seen.add(p)
        if p.is_file():
            return p
    return None


def document_preview_kind(doc: Document) -> str:
    path = resolve_document_path(doc)
    ext = path.suffix.lower() if path else Path(doc.filename).suffix.lower()
    if ext == ".pdf":
        return "pdf"
    if ext in (".png", ".jpg", ".jpeg", ".gif", ".webp", ".bmp"):
        return "image"
    if ext in (".md", ".txt", ".json", ".yaml", ".yml", ".csv"):
        return "text"
    if ext == ".docx":
        return "docx"
    return "download_only"


def document_preview_text(doc: Document, max_chars: int = 8000) -> str:
    path = resolve_document_path(doc)
    if not path:
        with get_session() as session:
            chunk = session.exec(
                select(DocumentChunk)
                .where(DocumentChunk.document_id == doc.id)
                .order_by(DocumentChunk.chunk_index)
                .limit(1)
            ).first()
        return (chunk.chunk_text or "")[:max_chars] if chunk else "Datei nicht auf dem Server gefunden."
    ext = path.suffix.lower()
    try:
        if ext == ".pdf":
            return extract_text_from_pdf(path)[:max_chars]
        if ext == ".docx":
            return extract_text_from_docx(path)[:max_chars]
        if ext in (".md", ".txt", ".json", ".yaml", ".yml", ".csv"):
            return path.read_text(encoding="utf-8", errors="replace")[:max_chars]
    except Exception:
        pass
    return "Vorschau konnte nicht gelesen werden."

def get_project_documents(project_key: str) -> List[Document]:
    """Gibt alle Dokumente zurück, die einem Projekt zugeordnet sind."""
    with get_session() as session:
        # Join ProjectDocumentLink -> Document
        statement = (
            select(Document)
            .join(ProjectDocumentLink, ProjectDocumentLink.document_id == Document.id)
            .where(ProjectDocumentLink.project_key == project_key)
            .where(Document.is_deleted == False)
        )
        return session.exec(statement).all()

def link_document_to_project(project_key: str, doc_id: int) -> bool:
    with get_session() as session:
        # Check if already linked
        existing = session.exec(
            select(ProjectDocumentLink)
            .where(ProjectDocumentLink.project_key == project_key)
            .where(ProjectDocumentLink.document_id == doc_id)
        ).first()
        
        if existing:
            return True # Already linked
            
        link = ProjectDocumentLink(project_key=project_key, document_id=doc_id)
        session.add(link)
        session.commit()
        return True

def unlink_document_from_project(project_key: str, doc_id: int) -> bool:
    with get_session() as session:
        link = session.exec(
            select(ProjectDocumentLink)
            .where(ProjectDocumentLink.project_key == project_key)
            .where(ProjectDocumentLink.document_id == doc_id)
        ).first()
        
        if link:
            session.delete(link)
            session.commit()
        return True


def sync_documents_to_chromadb() -> Tuple[int, int, str]:
    """
    Synchronisiert alle Dokumente und deren Chunks aus SQLite zu ChromaDB.
    
    Returns:
        Tuple (synced_chunks_count, projects_count, status_message)
    """
    from .m07_chroma import get_chroma_client, get_or_create_project_collection, add_chunks_to_collection
    from .m03_db import Project
    
    total_chunks_synced = 0
    projects_synced = set()
    errors = []
    
    try:
        with get_session() as session:
            projects = session.exec(select(Project).where(Project.is_deleted == False)).all()
            
            for project in projects:
                try:
                    project_docs = get_project_documents(project.key)
                    if not project_docs:
                        continue
                    
                    collection = get_or_create_project_collection(project.id, EMBEDDING_MODEL)
                    
                    for doc in project_docs:
                        chunks = session.exec(
                            select(DocumentChunk)
                            .where(DocumentChunk.document_id == doc.id)
                            .order_by(DocumentChunk.chunk_index)
                        ).all()
                        
                        if not chunks:
                            continue
                        
                        chunk_texts = [c.chunk_text for c in chunks]
                        chunk_embeddings = []
                        for c in chunks:
                            if c.embedding:
                                try:
                                    emb = json.loads(c.embedding)
                                    chunk_embeddings.append(emb)
                                except:
                                    chunk_embeddings.append(None)
                            else:
                                chunk_embeddings.append(None)
                        
                        add_chunks_to_collection(
                            collection,
                            document_id=doc.id,
                            chunks=chunk_texts,
                            embeddings=chunk_embeddings
                        )
                        
                        total_chunks_synced += len(chunks)
                    
                    projects_synced.add(project.key)
                
                except Exception as e:
                    errors.append(f"Projekt '{project.key}': {str(e)}")
        
        status = f"✅ {total_chunks_synced} Chunks in {len(projects_synced)} Projekten synchronisiert"
        if errors:
            status += f". Fehler: {'; '.join(errors[:3])}"
        
        return total_chunks_synced, len(projects_synced), status
    
    except Exception as e:
        return 0, 0, f"❌ ChromaDB Sync fehlgeschlagen: {str(e)}"
